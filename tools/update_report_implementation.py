from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


REPORT_PATH = Path(r"C:\Users\carlo\Desktop\Projeto\SI-AOM - Relatorio - Implementacao.docx")


def find_heading(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text and paragraph.style.name == "Heading 1":
            return paragraph
    raise RuntimeError(f"Heading not found: {text}")


def delete_between(doc: Document, start_paragraph, end_paragraph) -> None:
    body = doc._body._element
    children = list(body)
    start_idx = children.index(start_paragraph._p)
    end_idx = children.index(end_paragraph._p)
    for child in children[start_idx + 1 : end_idx]:
        body.remove(child)


def add_before(anchor, text: str = "", style: str = "Normal"):
    paragraph = anchor.insert_paragraph_before(text)
    try:
        paragraph.style = style
    except KeyError:
        paragraph.style = "Normal"
    return paragraph


def add_text(anchor, text: str):
    paragraph = add_before(anchor, text, "Normal")
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_heading(anchor, text: str, level: int = 2):
    style = f"Heading {level}"
    paragraph = add_before(anchor, text, style)
    paragraph.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_code(anchor, code: str):
    for raw_line in code.strip("\n").splitlines():
        paragraph = add_before(anchor, raw_line.rstrip(), "Code")
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(8)
    add_before(anchor, "", "Normal")


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def populate_implementation(anchor) -> None:
    add_heading(anchor, "Visão geral da implementação")
    add_text(
        anchor,
        "A implementação do SI-AOM foi organizada como uma aplicação web "
        "modular, com separação clara entre interface, API, regras de negócio, "
        "persistência e componente preditiva. A solução utiliza React no "
        "frontend, FastAPI no backend, SQLAlchemy para o acesso à base de dados, "
        "Pydantic para validação dos dados trocados pela API e scikit-learn para "
        "o módulo de Machine Learning. Esta divisão permitiu construir o sistema "
        "por domínios funcionais: clientes, projetos, orçamentos, materiais, "
        "operações, serviços, realizado, comparação, dashboard, autenticação e "
        "previsão de custos.",
    )
    add_text(
        anchor,
        "A lógica principal foi concentrada no backend, evitando que o frontend "
        "tivesse de conhecer regras internas de cálculo, permissões ou "
        "persistência. O frontend apresenta formulários, tabelas, estados de "
        "carregamento e mensagens de erro; a API valida os dados, aplica regras "
        "de autorização, calcula totais, regista alterações e devolve respostas "
        "normalizadas. Esta opção torna o sistema mais fácil de testar e reduz o "
        "risco de divergência entre ecrãs diferentes.",
    )
    add_text(
        anchor,
        "A estrutura física do projeto reflete esta separação. No backend, os "
        "modelos SQLAlchemy descrevem as tabelas, os schemas Pydantic descrevem "
        "a entrada e saída da API, os routers expõem os endpoints e os services "
        "guardam regras reutilizáveis. No frontend, os services encapsulam as "
        "chamadas HTTP, os hooks guardam estado transversal e os módulos React "
        "implementam os fluxos de trabalho de cada área.",
    )
    add_code(
        anchor,
        """
backend/app/
  models/      # tabelas e relações SQLAlchemy
  schemas/     # contratos de entrada e saída da API
  routers/     # endpoints FastAPI por domínio
  services/    # regras de negócio e cálculos reutilizáveis
  etl/ e ml/   # preparação de dados, treino e inferência

frontend/src/
  services/    # cliente HTTP e funções por recurso
  hooks/       # estado transversal, como autenticação
  modules/     # ecrãs de negócio: orçamentos, realizado, ML...
  components/  # componentes comuns de navegação e autenticação
""",
    )

    add_heading(anchor, "Arranque da API e configuração")
    add_text(
        anchor,
        "O ponto de entrada do backend é o ficheiro main.py. Nele é criada a "
        "instância FastAPI, aplicada a política CORS, registados os routers por "
        "módulo e inicializada a cache dos modelos de ML durante o ciclo de vida "
        "da aplicação. Esta cache evita carregar os ficheiros joblib a partir do "
        "disco em todas as previsões, reduzindo a latência das chamadas ao "
        "endpoint de previsão.",
    )
    add_code(
        anchor,
        """
@asynccontextmanager
async def lifespan(app: FastAPI):
    cache = MLModelCache()
    cache.load_all()
    app.state.ml_cache = cache
    try:
        yield
    finally:
        cache.clear()

app = FastAPI(title=settings.app_name, lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=settings.cors_origins)
app.include_router(orcamentos.router, prefix="/orcamentos")
app.include_router(realizado.router, prefix="/realizado")
app.include_router(ml.router, prefix="/ml")
""",
    )
    add_text(
        anchor,
        "As variáveis de configuração ficam centralizadas em config.py, "
        "permitindo adaptar a aplicação a ambientes distintos sem alterar o "
        "código. A ligação à base de dados, a chave JWT, o algoritmo de "
        "assinatura, a duração dos tokens e as origens CORS são lidos a partir "
        "do ambiente ou de valores por omissão controlados. Em produção, a "
        "existência de segredos fracos é tratada como erro de configuração.",
    )

    add_heading(anchor, "Modelo de dados e persistência")
    add_text(
        anchor,
        "A persistência foi implementada com SQLAlchemy ORM. Cada entidade de "
        "negócio possui uma classe própria e as relações entre tabelas são "
        "definidas por chaves estrangeiras. Nos detalhes de orçamento, foi "
        "adotado o conceito de snapshot: o preço ou custo unitário usado no "
        "momento da criação da linha fica guardado no orçamento. Assim, se o "
        "preço base de um material, operação ou serviço mudar no futuro, os "
        "orçamentos antigos continuam auditáveis e reproduzíveis.",
    )
    add_code(
        anchor,
        """
class DetalheMaterialOrcamento(Base):
    __tablename__ = "detalhe_material_orcamento"

    id_orcamento = mapped_column(
        ForeignKey("orcamento.id_orcamento", ondelete="CASCADE"),
        nullable=False,
    )
    id_material = mapped_column(
        ForeignKey("material.id_material", ondelete="RESTRICT"),
        nullable=False,
    )
    quantidade = mapped_column(Numeric(12, 2), nullable=False)
    preco_unitario_snapshot = mapped_column(Numeric(14, 4), nullable=False)
    custo_total = mapped_column(Numeric(14, 2), nullable=False)
""",
    )
    add_text(
        anchor,
        "A regra de eliminação em cascata foi usada apenas quando o detalhe "
        "depende totalmente do registo principal, como acontece nas linhas de um "
        "orçamento. Já entidades de catálogo, como materiais, operações e "
        "serviços, usam restrição de eliminação para evitar apagar dados ainda "
        "referenciados por orçamentos. Esta distinção protege a consistência "
        "histórica do sistema.",
    )

    add_heading(anchor, "Validação de dados e contratos da API")
    add_text(
        anchor,
        "Os schemas Pydantic funcionam como contratos formais entre frontend e "
        "backend. Para cada recurso existem schemas de criação, atualização e "
        "leitura, o que permite controlar que campos entram no sistema e que "
        "campos são devolvidos ao utilizador. Esta camada também evita que "
        "campos calculados, como totais de orçamento, sejam recebidos "
        "diretamente do cliente quando devem ser calculados pelo servidor.",
    )
    add_text(
        anchor,
        "A API segue um padrão consistente: recebe um payload validado pelo "
        "schema, obtém uma sessão de base de dados através da dependência "
        "get_db, aplica as regras do domínio e devolve um schema de resposta. "
        "Erros de validação estrutural são devolvidos automaticamente pelo "
        "FastAPI; erros de negócio, como entidades inexistentes ou permissões "
        "insuficientes, são convertidos em respostas HTTP explícitas.",
    )

    add_heading(anchor, "Autenticação, sessões e perfis de acesso")
    add_text(
        anchor,
        "A autenticação foi construída com email, password cifrada por bcrypt e "
        "tokens JWT. Durante o login, o backend procura o utilizador, valida a "
        "password, confirma se a conta está ativa e emite um token assinado com "
        "o identificador do utilizador e o respetivo perfil. O token é depois "
        "enviado pelo frontend no cabeçalho Authorization em todas as chamadas "
        "protegidas.",
    )
    add_code(
        anchor,
        """
def verificar_password(password: str, password_hash: str) -> bool:
    return bcrypt.checkpw(
        _password_bytes(password),
        password_hash.encode("utf-8"),
    )

def criar_token(data: dict) -> str:
    payload = data.copy()
    payload["exp"] = datetime.now(timezone.utc) + timedelta(
        minutes=settings.jwt_expire_minutes
    )
    return jwt.encode(
        payload,
        settings.jwt_secret,
        algorithm=settings.jwt_algorithm,
    )
""",
    )
    add_text(
        anchor,
        "As permissões são aplicadas com dependências FastAPI. A função "
        "require_roles recebe os perfis autorizados para uma rota, normaliza o "
        "perfil do utilizador autenticado e bloqueia a operação com HTTP 403 "
        "quando o perfil não pertence ao conjunto permitido. Esta abordagem "
        "torna a política de acesso visível diretamente no router de cada "
        "módulo.",
    )
    add_code(
        anchor,
        """
def require_roles(*allowed_roles: str):
    normalized_allowed = {_normalize_role(role) for role in allowed_roles}

    def _role_dependency(current_user=Depends(get_current_user)):
        perfil = _normalize_role(current_user.perfil)
        if perfil not in normalized_allowed:
            raise HTTPException(status_code=403, detail="Sem permissao")
        return current_user

    return _role_dependency
""",
    )
    add_text(
        anchor,
        "Com esta estratégia, o perfil administrador concentra operações de "
        "gestão, o orçamentista trabalha sobre clientes, projetos e orçamentos, "
        "o perfil de produção regista o realizado e o gestor consulta dados de "
        "controlo e comparação. A separação é implementada no backend e "
        "refletida no menu do frontend, evitando que a interface apresente "
        "opções que o utilizador não pode executar.",
    )

    add_heading(anchor, "Lógica de orçamentação")
    add_text(
        anchor,
        "O orçamento é o núcleo funcional do SI-AOM. A implementação separa o "
        "cabeçalho do orçamento das suas linhas de detalhe: materiais, "
        "operações e serviços. O cabeçalho guarda o projeto, estado, margem e "
        "totais consolidados; as linhas guardam quantidades, tempos, preços "
        "unitários em snapshot e custos calculados. Sempre que uma linha é "
        "criada, alterada ou removida, o backend recalcula os totais do "
        "orçamento numa função única de serviço.",
    )
    add_text(
        anchor,
        "O cálculo dos materiais combina quantidade, preço unitário e "
        "desperdício. Nas operações são considerados o tempo de execução, tempo "
        "de setup e custo hora. Nos serviços externos é aplicada a quantidade "
        "ao custo unitário contratado. O valor total do orçamento resulta da "
        "soma das três componentes e o preço de venda é recalculado quando "
        "existe margem definida.",
    )
    add_code(
        anchor,
        """
def recalcular_totais_orcamento(db: Session, id_orcamento: int) -> Orcamento:
    total_materiais = db.scalar(
        select(func.coalesce(func.sum(DetalheMaterialOrcamento.custo_total), 0))
        .where(DetalheMaterialOrcamento.id_orcamento == id_orcamento)
    )
    total_operacoes = db.scalar(
        select(func.coalesce(func.sum(DetalheOperacaoOrcamento.custo_total), 0))
        .where(DetalheOperacaoOrcamento.id_orcamento == id_orcamento)
    )
    total_servicos = db.scalar(
        select(func.coalesce(func.sum(DetalheServicoOrcamento.custo_total), 0))
        .where(DetalheServicoOrcamento.id_orcamento == id_orcamento)
    )
    orcamento.custo_total_orcado = _q2(
        total_materiais + total_operacoes + total_servicos
    )
""",
    )
    add_text(
        anchor,
        "A função usa Decimal e arredondamento financeiro para duas casas "
        "decimais, reduzindo erros de precisão típicos de valores monetários em "
        "ponto flutuante. Ao concentrar a consolidação num único serviço, todos "
        "os endpoints que manipulam linhas de orçamento obtêm o mesmo resultado "
        "e os testes conseguem validar a regra num ponto claro do código.",
    )

    add_heading(anchor, "Registo do realizado e comparação")
    add_text(
        anchor,
        "Depois de o projeto avançar para execução, o sistema permite registar "
        "custos reais de materiais, operações e serviços. Esta área foi "
        "separada do orçamento para preservar a diferença entre valor previsto "
        "e valor executado. O orçamento mantém a estimativa inicial, enquanto "
        "o realizado representa o consumo ou custo efetivamente observado.",
    )
    add_text(
        anchor,
        "A comparação agrega os valores orçados e realizados por componente e "
        "calcula desvios absolutos e percentuais. A partir desses desvios são "
        "gerados indicadores para o dashboard e para o ecrã de comparação, "
        "permitindo identificar rapidamente se o problema veio de materiais, "
        "operações, serviços ou do conjunto do projeto. Esta mesma informação "
        "serve de base histórica para alimentar o módulo preditivo.",
    )

    add_heading(anchor, "Frontend React e comunicação com a API")
    add_text(
        anchor,
        "O frontend foi implementado em React com Vite. A aplicação inicia em "
        "App.jsx, apresenta o fluxo de autenticação quando não existe sessão e, "
        "depois do login, disponibiliza o menu principal com os módulos "
        "permitidos ao perfil do utilizador. Cada módulo de negócio concentra a "
        "interface de uma área funcional, mas não contém detalhes de URL ou "
        "construção manual de pedidos HTTP; essa responsabilidade fica nos "
        "services.",
    )
    add_text(
        anchor,
        "O ficheiro apiClient.js centraliza a comunicação com a API. Ele junta a "
        "base URL configurada no ambiente, adiciona o cabeçalho Authorization "
        "quando existe token, serializa o corpo em JSON e transforma erros HTTP "
        "em mensagens legíveis. Desta forma, todos os módulos tratam respostas e "
        "erros da mesma maneira.",
    )
    add_code(
        anchor,
        """
export async function apiRequest(path, { method = 'GET', token, body } = {}) {
  const response = await fetch(toApiUrl(path), {
    method,
    headers: {
      'Content-Type': 'application/json',
      ...(token ? { Authorization: `Bearer ${token}` } : {}),
    },
    body: body ? JSON.stringify(body) : undefined,
  })

  const responseType = response.headers.get('content-type') || ''
  const isJson = responseType.includes('application/json')
  const payload = isJson ? await response.json() : null

  if (!response.ok) {
    let message = `Erro HTTP ${response.status}`
    if (payload?.detail) {
      message = typeof payload.detail === 'string'
        ? payload.detail
        : JSON.stringify(payload.detail)
    }
    throw new Error(message)
  }

  return payload
}
""",
    )
    add_text(
        anchor,
        "A sessão do utilizador é gerida pelo hook useAuth. O hook lê o token "
        "guardado, valida-o com o endpoint /auth/me, limpa a sessão quando o "
        "token deixa de ser válido e disponibiliza funções de login e logout à "
        "aplicação. Esta abstração evita duplicação de lógica em componentes "
        "visuais e garante que a aplicação reage de forma previsível a sessões "
        "expiradas.",
    )
    add_code(
        anchor,
        """
export function useAuth() {
  const [token, setToken] = useState(() => readToken())
  const [user, setUser] = useState(null)

  const loginWithPassword = async (email, password) => {
    const payload = await login(email, password)
    writeToken(payload.access_token)
    setToken(payload.access_token)
  }

  useEffect(() => {
    if (token) getCurrentUser(token).then(setUser).catch(logout)
  }, [token])

  return { token, user, loginWithPassword, logout }
}
""",
    )
    add_text(
        anchor,
        "Nos módulos de domínio, o padrão é carregar dados ao montar o ecrã, "
        "manter o formulário em estado local, enviar apenas o payload necessário "
        "ao service correspondente e atualizar a tabela depois da operação. O "
        "módulo de orçamentos acrescenta um painel de detalhes para materiais, "
        "operações e serviços; o módulo de realizado foca-se nos custos reais; "
        "o módulo de comparação apresenta os desvios; e o módulo de ML recolhe "
        "parâmetros técnicos para obter uma estimativa antes ou durante a "
        "elaboração do orçamento.",
    )

    add_heading(anchor, "Módulo de Machine Learning")
    add_text(
        anchor,
        "O módulo de Machine Learning foi implementado como componente de apoio "
        "à decisão, não como substituto do orçamentista. O objetivo é estimar "
        "custos prováveis a partir de características conhecidas antes da "
        "execução do projeto: tipologia, complexidade, peso total, área, número "
        "de peças, material principal, tratamento de superfície e lead time. O "
        "resultado devolve custos por componente, custo total, tempo previsto, "
        "qualidade do modelo, confiança e alertas quando os parâmetros estão "
        "fora da faixa observada no treino.",
    )
    add_heading(anchor, "Extração e preparação dos dados", 3)
    add_text(
        anchor,
        "A fase ETL transforma os dados relacionais da base de dados num dataset "
        "tabular. Cada linha corresponde a um orçamento com valores orçados e "
        "realizados agregados. Para evitar data leakage, as colunas que "
        "representam custos orçados são mantidas para auditoria e comparação, "
        "mas removidas durante o treino dos modelos que devem prever custos "
        "reais. Apenas orçamentos com custo orçado positivo e pelo menos um "
        "registo de realizado são elegíveis para treino.",
    )
    add_text(
        anchor,
        "A construção do dataset fica isolada em app/etl, permitindo regenerar "
        "o CSV a partir da base de dados sempre que existem novos projetos "
        "executados. Esta estratégia torna o treino reprodutível: a mesma "
        "consulta, com os mesmos filtros, produz a fonte de dados usada pelos "
        "modelos.",
    )
    add_heading(anchor, "Treino, métricas e persistência", 3)
    add_text(
        anchor,
        "O problema foi dividido em submodelos independentes para materiais, "
        "operações, serviços e horas. Cada submodelo usa um pipeline scikit-learn "
        "com pré-processamento e RandomForestRegressor. O pré-processamento "
        "imputa valores numéricos pela mediana, imputa categóricos pela moda e "
        "codifica categorias com OneHotEncoder configurado para aceitar valores "
        "desconhecidos sem falhar a previsão.",
    )
    add_code(
        anchor,
        """
preprocessor = ColumnTransformer([
    ("num", SimpleImputer(strategy="median"), numeric_cols),
    ("cat", Pipeline([
        ("imputer", SimpleImputer(strategy="most_frequent")),
        ("onehot", OneHotEncoder(handle_unknown="ignore")),
    ]), categorical_cols),
])

pipeline = Pipeline([
    ("preprocess", preprocessor),
    ("model", RandomForestRegressor(n_estimators=400, random_state=42)),
])
pipeline.fit(X_train, y_train)
""",
    )
    add_text(
        anchor,
        "A avaliação combina holdout, métricas como R2, MAE e RMSE, validação "
        "cruzada e comparação contra uma baseline simples baseada na mediana. "
        "Além do ficheiro model.joblib, cada treino grava metrics.json com as "
        "features utilizadas, qualidade, intervalos numéricos observados e "
        "valores categóricos conhecidos. Essas métricas são usadas tanto para "
        "explicar a qualidade da previsão como para preencher opções no "
        "frontend.",
    )
    add_heading(anchor, "Inferência integrada na API", 3)
    add_text(
        anchor,
        "Na previsão, o serviço valida os parâmetros obrigatórios, alinha as "
        "features recebidas com as features esperadas pelo modelo, carrega o "
        "modelo a partir da cache ou do disco, executa a predição por submodelo "
        "e agrega os resultados. Valores negativos são truncados para zero, "
        "pois custos e horas não podem ser negativos no domínio do problema.",
    )
    add_code(
        anchor,
        """
for sub_model in ORCAMENTO_SUB_MODELS:
    model = cache.get_model(sub_model) if cache else None
    if model is None:
        model = joblib.load(model_path)

    X = _align_features(model, features_used, pd.DataFrame([parametros]))
    valor = max(0.0, float(model.predict(X)[0]))
    resultados[sub_model.replace("orcamento_", "")] = valor

custo_total = (
    resultados.get("materiais", 0)
    + resultados.get("operacoes", 0)
    + resultados.get("servicos", 0)
)
""",
    )
    add_text(
        anchor,
        "O serviço também produz alertas de faixa de treino. Se um peso, área, "
        "número de peças ou lead time estiver fora dos valores observados no "
        "dataset, a confiança apresentada ao utilizador é reduzida e a resposta "
        "inclui uma mensagem de aviso. O mesmo acontece quando uma categoria "
        "não foi vista no treino. Esta lógica torna a previsão mais transparente "
        "e evita que o utilizador interprete uma estimativa fora de contexto "
        "como valor absoluto.",
    )
    add_text(
        anchor,
        "A previsão é persistida na tabela previsao_ml em modo fail-soft: se a "
        "gravação falhar, a resposta de previsão continua a ser devolvida ao "
        "utilizador e a transação é revertida apenas para o registo histórico. "
        "Esta escolha impede que uma falha secundária de auditoria bloqueie a "
        "funcionalidade principal.",
    )

    add_heading(anchor, "Integração entre orçamento, realizado e ML")
    add_text(
        anchor,
        "O desenho do SI-AOM liga três momentos do ciclo de vida: estimar, "
        "executar e aprender. Na fase de estimativa, o utilizador cria o "
        "orçamento através de linhas de materiais, operações e serviços, "
        "podendo consultar uma previsão de ML como referência. Na fase de "
        "execução, a produção regista os custos reais. Na fase de análise, o "
        "sistema compara orçado e realizado, alimenta indicadores de gestão e "
        "disponibiliza dados históricos para novo treino dos modelos.",
    )
    add_text(
        anchor,
        "Esta integração evita que o módulo de ML seja uma peça isolada. A "
        "qualidade das previsões depende da qualidade do realizado registado; "
        "por isso, os ecrãs operacionais e a estrutura da base de dados foram "
        "pensados para recolher dados que também têm valor analítico. O sistema "
        "passa a melhorar à medida que mais projetos são fechados com dados "
        "reais consistentes.",
    )

    add_heading(anchor, "Estratégia de testabilidade")
    add_text(
        anchor,
        "A organização da implementação facilita a criação de testes. Regras "
        "com impacto financeiro, como recalcular totais de orçamento, foram "
        "colocadas em serviços independentes. Endpoints FastAPI usam "
        "dependências substituíveis, permitindo testar fluxos completos com uma "
        "base de dados de teste. No frontend, a centralização do cliente HTTP e "
        "dos services permite simular respostas da API e validar fluxos sem "
        "depender de um servidor real.",
    )
    add_text(
        anchor,
        "Esta testabilidade é importante porque o SI-AOM combina permissões, "
        "cálculos, persistência e inferência preditiva. Ao manter contratos "
        "claros entre camadas, torna-se possível testar cada área de forma "
        "isolada e, em seguida, validar fluxos integrados como login, criação de "
        "orçamentos, registo do realizado, comparação de desvios e previsão por "
        "Machine Learning.",
    )

    add_heading(anchor, "Síntese da implementação")
    add_text(
        anchor,
        "Em síntese, o SI-AOM foi construído como um sistema integrado de apoio "
        "à orçamentação industrial. O backend concentra regras de negócio, "
        "segurança, cálculos e persistência; o frontend disponibiliza uma "
        "experiência orientada aos perfis de utilização; e o módulo de Machine "
        "Learning acrescenta uma camada preditiva baseada em dados históricos. "
        "A separação por camadas, o uso de snapshots nos custos, a validação por "
        "schemas, as permissões por perfil e a cache dos modelos de ML são as "
        "decisões técnicas que sustentam a robustez da solução.",
    )


def main() -> None:
    doc = Document(REPORT_PATH)
    implementation = find_heading(doc, "Implementação")
    tests = find_heading(doc, "Testes e Resultados")
    delete_between(doc, implementation, tests)
    populate_implementation(tests)
    enable_field_updates(doc)
    doc.save(REPORT_PATH)


if __name__ == "__main__":
    main()
